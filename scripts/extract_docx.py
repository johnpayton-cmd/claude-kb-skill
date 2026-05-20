"""Extract text from a DOCX file for KB summarization.

Usage:
  uv run --python 3.12 --with python-docx extract_docx.py <path>
  uv run --python 3.12 --with python-docx extract_docx.py <path> --headings-only
  uv run --python 3.12 --with python-docx extract_docx.py <path> --tables-only
  uv run --python 3.12 --with python-docx extract_docx.py <path> --max-rows 5

Workflow:
  1. Run with --headings-only first to see document structure (like reading a PDF TOC).
  2. Run without flags to get full content (paragraphs + tables).
  3. If tables are large, use --max-rows to control output size.
"""

import sys
import argparse
import docx


def extract_docx(path, tables_only=False, headings_only=False, max_rows=10):
    doc = docx.Document(path)

    if not tables_only:
        print("=== PARAGRAPHS ===")
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name
            if headings_only and "heading" not in style.lower():
                continue
            print(f"[{style}] {text}")

    if not headings_only:
        print(f"\n=== TABLES ({len(doc.tables)} total) ===")
        for i, table in enumerate(doc.tables):
            print(f"\n--- Table {i + 1} ---")
            row_count = 0
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(c[:200] for c in cells)
                if line.strip(" |"):
                    print(line)
                    row_count += 1
                    if row_count >= max_rows:
                        remaining = len(table.rows) - row_count
                        if remaining > 0:
                            print(f"  ... ({remaining} more rows)")
                        break


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extract text from a DOCX file for KB summarization."
    )
    parser.add_argument("path", help="Path to the DOCX file")
    parser.add_argument(
        "--headings-only",
        action="store_true",
        help="Output only headings (document structure overview)",
    )
    parser.add_argument(
        "--tables-only",
        action="store_true",
        help="Output only table content",
    )
    parser.add_argument(
        "--max-rows",
        type=int,
        default=10,
        help="Maximum rows to show per table (default: 10)",
    )
    args = parser.parse_args()

    extract_docx(
        args.path,
        tables_only=args.tables_only,
        headings_only=args.headings_only,
        max_rows=args.max_rows,
    )
